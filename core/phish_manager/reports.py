from io import BytesIO
from datetime import datetime
import os
import glob

from .models import Event


def _register_unicode_font():
    """DejaVu font varsa kaydet, yoksa None döndür (Helvetica'ya düşer)."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    _base = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.path.join(_base, 'fonts', 'DejaVuSans.ttf'),
        os.path.join(_base, '..', 'fonts', 'DejaVuSans.ttf'),
        os.path.join(_base, '..', '..', 'fonts', 'DejaVuSans.ttf'),
        'C:/Windows/Fonts/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/dejavu/DejaVuSans.ttf',
    ] + glob.glob('/nix/store/*/share/fonts/truetype/DejaVuSans.ttf')

    for path in candidates:
        path = os.path.normpath(path)
        if os.path.isfile(path):
            try:
                pdfmetrics.registerFont(TTFont('DejaVu', path))
                pdfmetrics.registerFont(TTFont('DejaVu-Bold', path))
                return 'DejaVu', 'DejaVu-Bold'
            except Exception:
                continue
    return 'Helvetica', 'Helvetica-Bold'


# ─── PDF ──────────────────────────────────────────────────────────────────────

def generate_pdf_report(campaign):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    font_normal, font_bold = _register_unicode_font()

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, spaceAfter=6, fontName=font_bold)
    sub_style   = ParagraphStyle('Sub',   parent=styles['Normal'], fontSize=10, textColor=colors.grey, fontName=font_normal)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=12, spaceBefore=14, spaceAfter=6, fontName=font_bold)

    events = Event.objects.filter(campaign=campaign).select_related('target')
    sent      = events.count()
    opened    = events.filter(was_opened=True).count()
    clicked   = events.filter(was_clicked=True).count()
    submitted = events.filter(was_submitted=True).count()
    click_rate = f"{round(clicked / sent * 100, 1) if sent > 0 else 0}%"

    elements = []

    # ── Başlık
    elements.append(Paragraph("Cyber Aware — Kampanya Raporu", title_style))
    elements.append(Paragraph(campaign.campaign_name, ParagraphStyle('CName', parent=styles['Normal'], fontSize=14, spaceAfter=4, fontName=font_bold)))
    elements.append(Paragraph(f"Rapor tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}", sub_style))
    elements.append(Spacer(1, 0.4*cm))

    # ── Kampanya bilgileri
    elements.append(Paragraph("Kampanya Bilgileri", section_style))
    info_data = [
        ["Alan", "Değer"],
        ["Durum", campaign.status],
        ["Şablon", campaign.template.template_name],
        ["Hedef Grup", campaign.target_group if campaign.target_group else "Manuel seçim"],
        ["Başlangıç",  campaign.start_date.strftime('%d.%m.%Y %H:%M')],
        ["Bitiş",      campaign.end_date.strftime('%d.%m.%Y %H:%M')],
    ]
    info_table = Table(info_data, colWidths=[5*cm, 12*cm])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
        ('FONTNAME',   (0, 0), (-1, 0), font_bold),
        ('FONTNAME',   (0, 1), (-1, -1), font_normal),
        ('FONTSIZE',   (0, 0), (-1, -1), 10),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f2f2f2')]),
        ('GRID',       (0, 0), (-1, -1), 0.5, colors.grey),
        ('LEFTPADDING',  (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING',   (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING',(0, 0), (-1, -1), 5),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.4*cm))

    # ── Özet istatistikler
    elements.append(Paragraph("Özet İstatistikler", section_style))
    summary_data = [
        ["Gönderilen", "Açılan", "Tıklayan", "Credential Giren", "Tıklama Oranı"],
        [str(sent), str(opened), str(clicked), str(submitted), click_rate],
    ]
    summary_table = Table(summary_data, colWidths=[3.4*cm]*5)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2980b9')),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
        ('FONTNAME',   (0, 0), (-1, 0), font_bold),
        ('ALIGN',      (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE',   (0, 0), (-1, -1), 10),
        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#d6eaf8')),
        ('FONTNAME',   (0, 1), (-1, 1), font_bold),
        ('FONTSIZE',   (0, 1), (-1, 1), 12),
        ('GRID',       (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING',   (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING',(0, 0), (-1, -1), 6),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 0.4*cm))

    # ── Hedef detayları
    elements.append(Paragraph("Hedef Detayları", section_style))
    target_data = [["Ad", "Soyad", "E-posta", "Grup", "Durum", "Açtı", "Tıkladı", "Girdi"]]
    for ev in events:
        target_data.append([
            ev.target.first_name,
            ev.target.last_name,
            ev.target.email,
            ev.target.group or "-",
            ev.status,
            "✓" if ev.was_opened    else "✗",
            "✓" if ev.was_clicked   else "✗",
            "✓" if ev.was_submitted else "✗",
        ])

    col_widths = [2.5*cm, 2.5*cm, 5*cm, 2.5*cm, 2*cm, 1.5*cm, 1.8*cm, 1.5*cm]
    target_table = Table(target_data, colWidths=col_widths)
    target_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
        ('FONTNAME',   (0, 0), (-1, 0), font_bold),
        ('FONTNAME',   (0, 1), (-1, -1), font_normal),
        ('FONTSIZE',   (0, 0), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f2f2f2')]),
        ('GRID',       (0, 0), (-1, -1), 0.5, colors.grey),
        ('ALIGN',      (5, 0), (-1, -1), 'CENTER'),
        ('LEFTPADDING',  (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING',   (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING',(0, 0), (-1, -1), 4),
    ]))
    elements.append(target_table)

    doc.build(elements)
    buffer.seek(0)
    return buffer


# ─── WORD ─────────────────────────────────────────────────────────────────────

def generate_word_report(campaign):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    events = Event.objects.filter(campaign=campaign).select_related('target')
    sent      = events.count()
    opened    = events.filter(was_opened=True).count()
    clicked   = events.filter(was_clicked=True).count()
    submitted = events.filter(was_submitted=True).count()
    click_rate = f"{round(clicked / sent * 100, 1) if sent > 0 else 0}%"

    doc = Document()

    # ── Sayfa kenar boşlukları
    section = doc.sections[0]
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    # ── Başlık
    title = doc.add_heading('Cyber Aware — Kampanya Raporu', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(campaign.campaign_name)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True

    date_p = doc.add_paragraph(f"Rapor tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    date_p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Kampanya bilgileri
    doc.add_heading('Kampanya Bilgileri', level=1)
    info_rows = [
        ("Durum",        campaign.status),
        ("Şablon",       campaign.template.template_name),
        ("Hedef Grup",   campaign.target_group if campaign.target_group else "Manuel seçim"),
        ("Başlangıç",    campaign.start_date.strftime('%d.%m.%Y %H:%M')),
        ("Bitiş",        campaign.end_date.strftime('%d.%m.%Y %H:%M')),
    ]
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    hdr = info_table.rows[0].cells
    hdr[0].text = "Alan"
    hdr[1].text = "Değer"
    _bold_row(hdr)
    _shade_cell(hdr[0], "2C3E50")
    _shade_cell(hdr[1], "2C3E50")
    _white_text(hdr[0])
    _white_text(hdr[1])
    for label, value in info_rows:
        row = info_table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph()

    # ── Özet
    doc.add_heading('Özet İstatistikler', level=1)
    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = 'Table Grid'
    headers = ["Gönderilen", "Açılan", "Tıklayan", "Credential Giren", "Tıklama Oranı"]
    values  = [str(sent), str(opened), str(clicked), str(submitted), click_rate]
    for i, (h, v) in enumerate(zip(headers, values)):
        hdr_cell = summary_table.rows[0].cells[i]
        val_cell = summary_table.rows[1].cells[i]
        hdr_cell.text = h
        val_cell.text = v
        _shade_cell(hdr_cell, "2980B9")
        _white_text(hdr_cell)
        _bold_cell(val_cell)
        _center_cell(hdr_cell)
        _center_cell(val_cell)

    doc.add_paragraph()

    # ── Hedef detayları
    doc.add_heading('Hedef Detayları', level=1)
    col_headers = ["Ad", "Soyad", "E-posta", "Grup", "Durum", "Açtı", "Tıkladı", "Girdi"]
    detail_table = doc.add_table(rows=1, cols=len(col_headers))
    detail_table.style = 'Table Grid'
    hdr_cells = detail_table.rows[0].cells
    for i, h in enumerate(col_headers):
        hdr_cells[i].text = h
        _shade_cell(hdr_cells[i], "2C3E50")
        _white_text(hdr_cells[i])
        _bold_cell(hdr_cells[i])
        _center_cell(hdr_cells[i])

    for ev in events:
        row = detail_table.add_row().cells
        row[0].text = ev.target.first_name
        row[1].text = ev.target.last_name
        row[2].text = ev.target.email
        row[3].text = ev.target.group or "-"
        row[4].text = ev.status
        row[5].text = "✓" if ev.was_opened    else "✗"
        row[6].text = "✓" if ev.was_clicked   else "✗"
        row[7].text = "✓" if ev.was_submitted else "✗"
        for cell in row[5:]:
            _center_cell(cell)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ─── Yardımcı fonksiyonlar ────────────────────────────────────────────────────

def _shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _white_text(cell):
    from docx.shared import RGBColor
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # If no runs yet, set on paragraph level via XML
    if not any(para.runs for para in cell.paragraphs):
        for para in cell.paragraphs:
            run = para.add_run(para.text)
            para.clear()
            para.add_run(cell.text).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _bold_row(cells):
    for cell in cells:
        _bold_cell(cell)


def _bold_cell(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True


def _center_cell(cell):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
